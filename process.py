import json
import sys
import time
import shutil
import logging
from pathlib import Path
from email import policy
from email.parser import BytesParser
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SCRIPT_DIR = Path(__file__).parent
TEMPLATE_PATH = SCRIPT_DIR / "template.docx"
EML_WAIT_TIMEOUT = 60   # секунд ждать EML файл
EML_WAIT_INTERVAL = 2   # интервал проверки в секундах

# Используем именованный логгер — не трогаем конфиг рутового логгера.
# Когда process.py импортирует daemon.py, используется его basicConfig.
# Когда process.py запускается напрямую — вывод идёт в stderr (дефолт).
logger = logging.getLogger(__name__)


def human_size(num_bytes):
    for unit in ["B", "KB", "MB", "GB"]:
        if num_bytes < 1024 or unit == "GB":
            return f"{num_bytes:.0f} {unit}" if unit == "B" else f"{num_bytes:.2f} {unit}"
        num_bytes /= 1024


def parse_eml(eml_path):
    with open(eml_path, "rb") as f:
        msg = BytesParser(policy=policy.default).parse(f)

    text_plain_parts = []
    text_html_parts = []

    for part in msg.walk():
        if part.is_multipart():
            continue
        ctype = part.get_content_type()
        disp = part.get_content_disposition()
        filename = part.get_filename()

        if disp == "attachment" or filename:
            continue

        if ctype == "text/plain":
            try:
                text_plain_parts.append(part.get_content())
            except Exception:
                pass
        elif ctype == "text/html":
            try:
                text_html_parts.append(part.get_content())
            except Exception:
                pass

    body_text = ""
    if text_plain_parts:
        body_text = "\n\n".join(t.strip() for t in text_plain_parts if t and t.strip())
    elif text_html_parts:
        body_text = "\n\n".join(t.strip() for t in text_html_parts if t and t.strip())

    attachments = []
    for part in msg.walk():
        if part.is_multipart():
            continue
        disp = part.get_content_disposition()
        filename = part.get_filename()
        if disp != "attachment" and not filename:
            continue
        filename = filename or "attachment.bin"
        data = part.get_payload(decode=True)
        size = len(data) if data else 0
        attachments.append((filename, size))

    return body_text, attachments


def parse_json(json_path):
    with open(json_path, encoding="utf-8") as f:
        data = json.load(f)

    record = data["data"]["mail"][0]
    meta = record["meta"]

    receiver_raw = meta.get("receiver", [])
    if isinstance(receiver_raw, list):
        receiver_val = "\n".join(str(v) for v in receiver_raw)
    else:
        receiver_val = str(receiver_raw)

    eml_filename = record.get("filename") or meta.get("filename") or ""
    if isinstance(eml_filename, list):
        eml_filename = eml_filename[0] if eml_filename else ""

    return {
        "start_time":   str(record.get("start_time", "")),
        "stop_time":    str(record.get("stop_time", "")),
        "sender":       str(meta.get("sender", "")),
        "receiver":     receiver_val,
        "eml_filename": eml_filename,
        "order_code":   str(record.get("order_code", "")).strip(),
    }


def clear_cell(cell):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    for para in cell.paragraphs[1:]:
        para._element.getparent().remove(para._element)


def replace_cell_text(cell, new_text):
    import copy
    lines = str(new_text).split("\n")
    clear_cell(cell)
    first_para = cell.paragraphs[0]

    # Сохранить rPr первого рана для единообразного шрифта во всех строках
    rPr_template = None
    if first_para.runs:
        existing_rPr = first_para.runs[0]._element.find(qn("w:rPr"))
        if existing_rPr is not None:
            rPr_template = copy.deepcopy(existing_rPr)
        first_para.runs[0].text = lines[0]
    else:
        first_para.add_run(lines[0])

    for line in lines[1:]:
        new_para = OxmlElement("w:p")
        # Убрать отступы между строками
        pPr = OxmlElement("w:pPr")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        pPr.append(spacing)
        new_para.append(pPr)
        # Ран с тем же форматированием что и первая строка
        r = OxmlElement("w:r")
        if rPr_template is not None:
            r.append(copy.deepcopy(rPr_template))
        t = OxmlElement("w:t")
        t.text = line
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        new_para.append(r)
        cell._element.append(new_para)


def fill_table_cell(table, label, value):
    for row in table.rows:
        unique_cells = []
        seen = set()
        for cell in row.cells:
            if id(cell) not in seen:
                seen.add(id(cell))
                unique_cells.append(cell)
        for i, cell in enumerate(unique_cells):
            if cell.text.strip() == label and i + 1 < len(unique_cells):
                replace_cell_text(unique_cells[i + 1], value)
                return


def fill_attachments_cell(table, order_code, attachments):
    """Найти ячейку с 'test', вставить вложения сверху, фиксированный текст не трогать."""
    import copy
    prefix = order_code if order_code else "нет"
    lines = [f"{prefix} {name} qwerty {size} байт" for name, size in attachments] if attachments else [prefix]

    for row in table.rows:
        seen = set()
        for cell in row.cells:
            if id(cell) in seen:
                continue
            seen.add(id(cell))
            if "test" not in cell.text:
                continue

            # Сохранить rPr из placeholder-рана для единообразного шрифта
            rPr_template = None
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text.strip() == "test":
                        existing_rPr = run._element.find(qn("w:rPr"))
                        if existing_rPr is not None and rPr_template is None:
                            rPr_template = copy.deepcopy(existing_rPr)

            # Удалить только placeholder-раны "test", фиксированный текст оставить
            for para in list(cell.paragraphs):
                for run in list(para.runs):
                    if run.text.strip() == "test":
                        run._element.getparent().remove(run._element)
                # Удалить параграф если стал пустым (но не единственный)
                if not para._element.findall(".//" + qn("w:t")) and len(cell.paragraphs) > 1:
                    para._element.getparent().remove(para._element)

            # Вставить строки вложений в начало ячейки
            first_para = cell.paragraphs[0]
            insert_idx = list(cell._element).index(first_para._element)
            for line in reversed(lines):
                new_para = OxmlElement("w:p")
                pPr = OxmlElement("w:pPr")
                spacing = OxmlElement("w:spacing")
                spacing.set(qn("w:before"), "0")
                spacing.set(qn("w:after"), "0")
                pPr.append(spacing)
                new_para.append(pPr)
                r = OxmlElement("w:r")
                if rPr_template is not None:
                    r.append(copy.deepcopy(rPr_template))
                t = OxmlElement("w:t")
                t.text = line
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                r.append(t)
                new_para.append(r)
                cell._element.insert(insert_idx, new_para)
            return


def make_paragraph(text, bold=False, compact=False):
    p = OxmlElement("w:p")
    if compact:
        pPr = OxmlElement("w:pPr")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        pPr.append(spacing)
        p.append(pPr)
    if text:
        r = OxmlElement("w:r")
        if bold:
            rPr = OxmlElement("w:rPr")
            rPr.append(OxmlElement("w:b"))
            r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        p.append(r)
    return p


def insert_text_after_table(doc, body_text):
    """Вставить текст сообщения сразу после таблицы с данными."""
    if not body_text:
        return

    body = doc.element.body
    all_tables = body.findall(qn("w:tbl"))
    if not all_tables:
        return

    data_table = all_tables[1] if len(all_tables) > 1 else all_tables[0]
    children = list(body)
    idx = children.index(data_table)

    # Удалить пустые параграфы сразу после таблицы данных
    for child in children[idx + 1:]:
        if child.tag == qn("w:p") and child.find(".//" + qn("w:t")) is None:
            body.remove(child)
        else:
            break

    # Позиция вставки — сразу после таблицы данных
    idx = list(body).index(data_table) + 1

    # Небольшой отступ после таблицы
    inserts = [make_paragraph("", compact=False)]
    for line in body_text.split("\n"):
        inserts.append(make_paragraph(line, compact=True))

    for para in reversed(inserts):
        body.insert(idx, para)


def process_json(json_file):
    json_path = Path(json_file)
    watch_dir = json_path.parent

    logger.info(f"Processing: {json_path}")

    try:
        data = parse_json(json_path)
    except Exception as e:
        logger.error(f"Failed to parse JSON {json_path}: {e}")
        return

    body_text = ""
    attachments = []
    eml_filename = data["eml_filename"]
    if eml_filename:
        eml_path = watch_dir / eml_filename
        if not eml_path.exists():
            logger.info(f"Waiting for EML: {eml_path}")
            elapsed = 0
            while not eml_path.exists() and elapsed < EML_WAIT_TIMEOUT:
                time.sleep(EML_WAIT_INTERVAL)
                elapsed += EML_WAIT_INTERVAL
        if eml_path.exists():
            try:
                body_text, attachments = parse_eml(eml_path)
            except Exception as e:
                logger.error(f"Failed to parse EML {eml_path}: {e}")
        else:
            logger.warning(f"EML not found after {EML_WAIT_TIMEOUT}s, skipping: {eml_path}")

    try:
        doc = Document(TEMPLATE_PATH)
    except Exception as e:
        logger.error(f"Failed to open template: {e}")
        return

    tables = doc.tables
    header_table = tables[0]
    data_table = tables[1] if len(tables) > 1 else tables[0]

    # Таблица 0 (шапка): ячейка "test" — вложения из EML
    fill_attachments_cell(header_table, data["order_code"], attachments)

    # Таблица 1 (данные): время, отправитель, получатель
    fill_table_cell(data_table, "Бошланиш вақти:", data["start_time"])
    fill_table_cell(data_table, "Тугаш вақти:", data["stop_time"])
    fill_table_cell(data_table, "Ким:", data["sender"])
    fill_table_cell(data_table, "Кимга:", data["receiver"])

    # Текст сообщения — под таблицей данных
    insert_text_after_table(doc, body_text)

    output_path = watch_dir / (json_path.stem + ".docx")
    tmp_path = output_path.with_suffix(".tmp")
    try:
        doc.save(tmp_path)
        shutil.move(tmp_path, output_path)  # атомарная операция — либо целый файл либо ничего
        logger.info(f"Saved: {output_path}")
    except Exception as e:
        logger.error(f"Failed to save {output_path}: {e}")
        tmp_path.unlink(missing_ok=True)  # убрать мусор если что-то пошло не так


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python process.py file.json")
        sys.exit(1)
    # При запуске напрямую — вывод логов в консоль
    logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
    process_json(sys.argv[1])
